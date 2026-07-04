"""Assessment Prep Binder — a pure docx generator.

Named "Assessment Prep Binder" deliberately: it is material you bring TO an
assessment, not proof you passed one. It reuses the SAME dashboard_summary the UI
shows, so the exported verdict can never disagree with the screen.

PURE: takes plain data, returns docx bytes. No Streamlit, no disk I/O, no network.
`today` is a parameter. Testable by parsing the returned document.
"""

import io

from docx import Document

import disclosures
from logic.readiness import (
    conditional_eligibility, dashboard_summary, evidence_index_from_register,
    poam_eligible,
)
from logic.scoring import IMPLEMENTED, NA_NOT_PERMITTED, NOT_IMPLEMENTED

_STATUS_LABEL = {
    IMPLEMENTED: "Implemented",
    NOT_IMPLEMENTED: "Not implemented",
    "partial_alt": "Partial",
    NA_NOT_PERMITTED: "N/A (not permitted)",
}


def build_binder(*, org_name, catalog, assessment, rules, evidence=None,
                 poam=None, org_profile=None, today=None) -> bytes:
    """Return the Assessment Prep Binder as .docx bytes."""
    evidence = evidence or {}
    poam = poam or {}
    org_profile = org_profile or {}
    by_id = {c["id"]: c for c in catalog}
    ev_index = evidence_index_from_register(evidence)
    summary = dashboard_summary(assessment, catalog, rules, ev_index, today=today)
    elig = conditional_eligibility(assessment, catalog, rules)

    doc = Document()
    org_label = org_name or "Untitled organization"

    # ---- cover ----
    doc.add_heading("CMMC Level 2 Assessment Prep Binder", level=0)
    doc.add_paragraph(f"Organization: {org_label}")
    doc.add_paragraph(f"SPRS score (self-estimate): {summary.score}")
    verdict = ("Projected: Final-ready" if summary.final_ready
               else "Projected: Conditional-eligible" if summary.conditional_eligible
               else "Not conditionally ready")
    doc.add_paragraph(f"Projected readiness (self-estimate): {verdict}")
    doc.add_paragraph(disclosures.READINESS_QUALIFIER)
    doc.add_paragraph(disclosures.DISCLAIMER)

    # ---- app -> CMMC status mapping ----
    doc.add_heading("What this readiness means vs. official CMMC status", level=1)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    h = t.rows[0].cells
    h[0].text, h[1].text, h[2].text = (
        "This binder shows (self-estimate)", "Official CMMC meaning", "How it is actually earned")
    for a, b, c in disclosures.APP_TO_CMMC_STATUS:
        r = t.add_row().cells
        r[0].text, r[1].text, r[2].text = a, b, c
    doc.add_paragraph("This tool and this binder confer no CMMC status.")

    # ---- executive summary (compound gate + SSP validity) ----
    doc.add_heading("Executive summary", level=1)
    if summary.ssp_missing:
        doc.add_paragraph(
            "SSP VALIDITY GATE: No System Security Plan (3.12.4). Under the DoD "
            "methodology the assessment cannot be completed — there is no valid SPRS "
            "score. This binder is for planning only.")
    doc.add_paragraph(
        f"Reaching 88 is necessary but not sufficient. Conditional-eligible: "
        f"{'YES' if summary.conditional_eligible else 'NO'}.")
    for reason in elig.reasons:
        doc.add_paragraph(reason, style="List Bullet")
    doc.add_paragraph(
        f"Open requirements — 5-pt: {summary.open_5pt}, 3-pt: {summary.open_3pt}, "
        f"1-pt: {summary.open_1pt}. Evidence register coverage: "
        f"{summary.evidence_covered}/{summary.evidence_applicable}.")
    if elig.blocking_ids:
        doc.add_paragraph("Blockers that cannot be deferred to a POA&M (must be MET):")
        for cid in elig.blocking_ids:
            doc.add_paragraph(f"{cid} — {by_id[cid]['short_title']}", style="List Bullet")

    # ---- SSP skeleton ----
    doc.add_heading("System Security Plan (skeleton — to be completed)", level=1)
    doc.add_paragraph(f"Legal name: {org_profile.get('legal_name', org_label)}")
    doc.add_paragraph(f"CAGE code: {org_profile.get('cage_code', '________')}")
    doc.add_paragraph(f"System boundary: {org_profile.get('system_boundary', '________')}")
    doc.add_paragraph("Attach the asset inventory and a network diagram (this tool "
                      "does not generate a diagram).")

    # ---- control-implementation matrix (all 110) ----
    doc.add_heading("Control implementation matrix", level=1)
    mt = doc.add_table(rows=1, cols=5)
    mt.style = "Light Grid Accent 1"
    hdr = mt.rows[0].cells
    for i, name in enumerate(("Control", "Requirement", "Weight", "Status", "Evidence")):
        hdr[i].text = name
    for c in catalog:
        cid = c["id"]
        status = assessment.get(cid, NOT_IMPLEMENTED)
        w = "SSP" if c["weight"] == "NA" else str(c["weight"])
        n_ev = len(evidence.get(cid, []))
        cov = "operational+reviewed" if ev_index.get(cid, {}).get("has_operational_final") \
            else (f"{n_ev} entr{'y' if n_ev == 1 else 'ies'}" if n_ev else "—")
        row = mt.add_row().cells
        row[0].text = cid
        row[1].text = c["short_title"]
        row[2].text = w
        row[3].text = _STATUS_LABEL.get(status, status)
        row[4].text = cov

    # ---- POA&M ----
    doc.add_heading("Plan of Action & Milestones", level=1)
    if not poam:
        doc.add_paragraph("No POA&M items recorded.")
    else:
        pt = doc.add_table(rows=1, cols=4)
        pt.style = "Light Grid Accent 1"
        ph = pt.rows[0].cells
        for i, name in enumerate(("Control", "Requirement", "POA&M-eligible?", "Target date")):
            ph[i].text = name
        for cid, p in sorted(poam.items()):
            c = by_id.get(cid, {"short_title": "?"})
            status = assessment.get(cid, NOT_IMPLEMENTED)
            elig_flag = "Yes" if (cid in by_id and poam_eligible(by_id[cid], status, rules)) \
                else "NO — must be MET"
            row = pt.add_row().cells
            row[0].text = cid
            row[1].text = c["short_title"]
            row[2].text = elig_flag
            row[3].text = p.get("target_date", "")

    # ---- evidence register index (metadata only) ----
    doc.add_heading("Evidence register index (metadata only — no files)", level=1)
    any_ev = any(evidence.values())
    if not any_ev:
        doc.add_paragraph("No evidence register entries recorded.")
    else:
        et = doc.add_table(rows=1, cols=6)
        et.style = "Light Grid Accent 1"
        eh = et.rows[0].cells
        for i, name in enumerate(("Control", "Title", "Owner", "Location", "Doc", "Operational/Review")):
            eh[i].text = name
        for cid in sorted(evidence, key=lambda x: tuple(int(p) for p in x.split("."))):
            for e in evidence[cid]:
                row = et.add_row().cells
                row[0].text = cid
                row[1].text = e.get("title", "")
                row[2].text = e.get("owner", "")
                row[3].text = e.get("location_uri", "")
                row[4].text = e.get("doc_status", "")
                row[5].text = f"{e.get('impl_status', '')}/{e.get('review_status', '')}"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
