"""Tests for the pure Assessment Prep Binder generator."""

import io
import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document

from export.binder import build_binder
from logic.catalog import controls, load_sample, poam_rules

CAT = controls()
RULES = poam_rules()


def _doc(data):
    return Document(io.BytesIO(data))


class TestBinder(unittest.TestCase):
    def setUp(self):
        self.sample = load_sample()
        self.data = build_binder(
            org_name="Sample Shop", catalog=CAT, assessment=self.sample["statuses"],
            rules=RULES, evidence=self.sample["evidence"], poam=self.sample["poam"],
        )

    def test_returns_bytes(self):
        self.assertIsInstance(self.data, bytes)
        self.assertGreater(len(self.data), 0)

    def test_matrix_has_all_110_controls(self):
        doc = _doc(self.data)
        # the control-implementation matrix is the table with a 110-row body
        big = [t for t in doc.tables if len(t.rows) == 111]  # header + 110
        self.assertEqual(len(big), 1, "expected exactly one 110-row control matrix")

    def test_mapping_table_and_language(self):
        text = "\n".join(p.text for p in _doc(self.data).paragraphs)
        self.assertIn("Assessment Prep Binder", text)
        self.assertIn("self-estimate", text)
        self.assertIn("confer", text.lower())
        # honesty: must not claim certification/compliance
        low = text.lower()
        for forbidden in ("certified", "is compliant", "guaranteed", "audit-approved"):
            self.assertNotIn(forbidden, low)

    def test_sample_binder_says_not_conditionally_ready(self):
        text = "\n".join(p.text for p in _doc(self.data).paragraphs)
        self.assertIn("Not conditionally ready", text)
        # the four blockers are named
        for cid in ("3.4.8", "3.13.6", "3.14.6", "3.1.20"):
            self.assertIn(cid, text)

    def test_no_file_bytes_only_metadata(self):
        # evidence index lists titles/owners/locations, never file content
        doc = _doc(self.data)
        joined = "\n".join(
            c.text for t in doc.tables for r in t.rows for c in r.cells)
        self.assertIn("Badge access policy + Q2 facility log", joined)  # a title
        # a signed-but-not-operational entry must NOT count as covered
        self.assertNotIn("operational+reviewed", _cell_for(doc, "3.4.8"))


def _cell_for(doc, cid):
    for t in doc.tables:
        for r in t.rows:
            if r.cells and r.cells[0].text == cid:
                return " | ".join(c.text for c in r.cells)
    return ""


if __name__ == "__main__":
    unittest.main()
